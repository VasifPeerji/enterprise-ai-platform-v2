import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
except ImportError:
    print("python-docx is not installed. Installing it now...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn

def set_style_font(style, name, size, bold=False, italic=False, underline=False):
    font = style.font
    font.name = name
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic
    font.underline = underline
    font.color.rgb = RGBColor(0, 0, 0)
    
def add_h1(doc, text):
    p = doc.add_paragraph(text, style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_h2(doc, text):
    return doc.add_paragraph(text, style='Heading 2')

def add_h3(doc, text):
    return doc.add_paragraph(text, style='Heading 3')

def add_p(doc, text):
    p = doc.add_paragraph(text, style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def main():
    doc = Document()

    # --- STYLE CONFIGURATION ---
    style_normal = doc.styles['Normal']
    set_style_font(style_normal, 'Times New Roman', 12)
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style_normal.paragraph_format.line_spacing = 1.15

    style_h1 = doc.styles['Heading 1']
    set_style_font(style_h1, 'Times New Roman', 20, bold=True, italic=True, underline=True)
    style_h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    style_h2 = doc.styles['Heading 2']
    set_style_font(style_h2, 'Times New Roman', 16, bold=True)

    style_h3 = doc.styles['Heading 3']
    set_style_font(style_h3, 'Times New Roman', 12, bold=True)

    # --- DOCUMENT CONTENT GENERATION ---
    
    # 1. Title Page / Cover
    for _ in range(5): add_p(doc, "")
    add_h1(doc, "SYNOPSIS")
    for _ in range(2): add_p(doc, "")
    add_h2(doc, "Enterprise AI Platform: A 7-Layer Cognitive and Transactional Architecture").alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(15): add_p(doc, "")
    add_p(doc, "Prepared By: Enterprise Architecture Team").alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_p(doc, "Date: February 2026").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 2. Problem Context
    add_h1(doc, "1. Problem Context")
    add_h2(doc, "1.1 The Shift in Enterprise AI Landscape")
    add_p(doc, "In recent years, the integration of Artificial Intelligence (AI), particularly Large Language Models (LLMs), into enterprise workflows has moved from an experimental phase to an imperative operational strategy. The ability to parse unstructured data, generate code, and answer domain-specific questions has unlocked unprecedented productivity gains. Enterprises are rapidly shifting from simple chatbot proofs-of-concept to deeply integrated 'Agentic' workflows capable of managing multi-step tasks across isolated departments.")
    add_p(doc, "However, early adopters have encountered significant friction when attempting to scale these solutions across hundreds of internal and external use cases. Connecting generic AI capabilities directly to sensitive backend systems introduces unprecedented challenges around access control, token leakage, and unpredictable latency variations. What works efficiently for a simple code-generation task completely misfires when asked to validate financial transactions.")

    add_h2(doc, "1.2 Core Challenges Addressed")
    add_h3(doc, "1.2.1 Vendor Lock-in and Model Obsolescence")
    add_p(doc, "The rapid pace of AI advancement implies that today's state-of-the-art model will become obsolete in months. Enterprises building monolithic applications directly coupled to a single provider (e.g., OpenAI, Anthropic, Google) face massive refactoring costs every time they wish to switch providers. Furthermore, relying entirely on frontier models for trivial inferences incurs unsustainable long-term API expenditures.")
    
    add_h3(doc, "1.2.2 The Hallucination and Action Risk")
    add_p(doc, "While AI models are proficient at text generation, their stochastic nature introduces hallucinations. When AI systems are endowed with 'agency'—the ability to execute APIs, modify databases, or send emails—the impact of a hallucinated parameter can be catastrophic. Without a strict separation between \"thinking\" operations and \"doing\" operations, enterprises expose themselves to vast regulatory liabilities.")

    add_h3(doc, "1.2.3 Lack of Multi-Tenancy and Governance")
    add_p(doc, "Enterprise systems must support diverse tenants—different departments, external clients, or restricted environments. Current out-of-the-box abstractions fail to strictly isolate memory vectors across departments, leading to scenarios where an HR agent might inadvertently expose salary knowledge when querying an improperly partitioned unified vector database.")
    
    add_h3(doc, "1.2.4 Evaluation and CI/CD for AI")
    add_p(doc, "Traditional software engineering relies on deterministic tests. AI applications, however, yield non-deterministic outputs. Establishing When a system update is pushed, or a core system prompt is mutated, traditional unit tests cannot definitively prove that the accuracy of a specialized response remains intact.")
    doc.add_page_break()

    # 3. Proposed System
    add_h1(doc, "2. Proposed System: 7-Layer Architecture")
    add_p(doc, "To resolve the aforementioned challenges, we propose a strict, decoupled, 7-Layer Enterprise AI Architecture. By dividing responsibilities across distinct operational layers, the platform guarantees scalability, security, and dynamic model flexibility. This enforces a unidirectional flow where higher layers manipulate business abstractions, supported strictly by lower engine capabilities.")

    add_h2(doc, "2.1 Layer 0: Model & Multimodal Infrastructure")
    add_p(doc, "Layer 0 serves as the foundational infrastructure that abstracts all physical model interactions. This prevents higher layers from being tightly bound to specific APIs. Through a dynamic unified interface (LiteLLM), cognitive loads are programmatically separated between heavyweight frontier endpoints and local open-weight inference runners depending entirely on the contextual classification. ")
    add_h3(doc, "Key Components in Layer 0:")
    add_p(doc, "• Model Registry: A dynamic catalog of available AI models.\n• Dynamic Router: Intelligently routes requests based on latency, cost, and complexity.\n• Determinism Controller: Adjusts model temperature and output probability to ensure regulatory compliance.")
    
    add_h2(doc, "2.2 Layer 1: Core Intelligence (Cognitive Brain)")
    add_p(doc, "Layer 1 functions as the pure computational intelligence layer. It is explicitly restricted to read-only operations to ensure zero destructive side-effects during planning, reasoning, or content generation. It handles all forms of unstructured comprehension, assembling context through distributed retrieval networks.")
    add_p(doc, "This layer manages Intent Routing, determining if the user request is purely conversational, requires RAG (Retrieval-Augmented Generation), or involves transactional intent. It contains the semantic memory abstractions responsible for injecting conversational history continuously into active memory states.")

    add_h2(doc, "2.3 Layer 2: Transaction & Agent Runtime (Doing Brain)")
    add_p(doc, "Unlike Layer 1, Layer 2 is the ONLY system layer authorized to perform side-effects. This division is the bedrock of the entire architectural security premise.")
    add_p(doc, "It leverages a stateful graph-based workflow orchestrator to manage multi-step agent actions. It introduces mechanisms like Idempotency Keys (to prevent double execution of identical tasks across distributed nodes) and immense Policy & Risk Evaluation protocols that halt and request explicit Human-In-The-Loop approval before submitting modifications.")

    add_h2(doc, "2.4 Layer 3: Domain Engine (Reusability)")
    add_p(doc, "Layer 3 bridges generic AI capabilities with specific business realities. While generic AI can write python, it does not implicitly understand the metadata schema of an internal company database without rigorous contextual injection.")
    add_h3(doc, "Key Mechanisms:")
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Role'
    row_cells = table.add_row().cells
    row_cells[0].text = 'Ontology Mapper'
    row_cells[1].text = 'Aligns generic LLM output to strictly defined domain schemas.'
    row_cells = table.add_row().cells
    row_cells[0].text = 'Universal Data Ingestion'
    row_cells[1].text = 'ETL pipelines designed to chunk, embed, and index enterprise documents into Qdrant vector databases.'

    add_h2(doc, "2.5 Layer 4: Platform Engine (Multi-Tenancy)")
    add_p(doc, "The Platform Engine handles multi-tenancy logic dynamically at the HTTP access layer. By parsing incoming JWTs and custom tenant headers, it ensures that isolation exists at the database, cognitive memory, and prompt levels, routing traffic exclusively against partitioned datasets.")

    add_h2(doc, "2.6 Layer 5: Governance, Observability & Learning")
    add_p(doc, "Comprehensive observability forms the backbone of enterprise trust. Every interaction is traced, measuring the latency across individual layers, RAG grounding, and cost analytics per token, funneled directly into centralized observability endpoints for anomaly detection.")

    add_h2(doc, "2.7 Layer 6: AI Ops & Evaluation")
    add_p(doc, "The highest layer acts as a gatekeeper for production releases. Using 'Golden Datasets', it runs thousands of deterministic and LLM-as-a-Judge evaluations to trap regressions in accuracy, effectively replacing legacy CI/CD pipelines with AI-specific continuous evaluation schemas.")
    doc.add_page_break()

    # 4. Objectives up to Deployment Level
    add_h1(doc, "3. Clear and Well-Defined Objectives")
    
    add_h2(doc, "3.1 Technical Objectives")
    add_p(doc, "1. Implement zero-trust boundaries between Cognitive (Layer 1) and Transactional (Layer 2) operations ensuring complete computational isolation.")
    add_p(doc, "2. Achieve strict latency budgets: <200ms for Fast Triage routes, <800ms for generic routing queries, masking any complex asynchronous verification logic within streamable UI endpoints.")
    add_p(doc, "3. Ensure 100% decoupling from monolithic LLM providers, utilizing an abstract unified proxy interface that permits hot-swapping baseline reasoning models without requiring extensive monolithic code base modifications.")

    add_h2(doc, "3.2 Operational and Business Objectives")
    add_p(doc, "1. Reduce active generation token spend drastically through intelligent prompt caching, explicit compression techniques in RAG pipelines, and dynamic tiering (e.g., routing structurally trivial queries to smaller fine-tuned local models). ")
    add_p(doc, "2. Onboard new corporate tenants in under 15 minutes with distinct organizational namespaces, deploying dedicated vector isolation clustering, and strictly federated role-based access control. ")
    
    add_h2(doc, "3.3 Deployment Lifecycle Objectives")
    add_p(doc, "1. Establish a multi-environment CI/CD pipeline featuring Continuous Synthetic Evaluation.")
    add_p(doc, "2. Implement Zero-Downtime rollouts via Kubernetes orchestration, ensuring routing topology traffic is incrementally shifted via Canary evaluations. Reversion macros must execute in sub-minute intervals if observability metrics breach specific degradation thresholds.")
    doc.add_page_break()

    # 5. Methodology
    add_h1(doc, "4. Methodology")
    add_p(doc, "The system development lifecycle (SDLC) is anchored in Agile principles paired with specific Machine Learning Operations (MLOps) adaptations tailored explicitly for generative non-deterministic outputs.")

    add_h2(doc, "4.1 Architectural Paradigms")
    add_p(doc, "The project strictly adheres to a domain-driven design philosophy. Dependencies between layers enforce a unidirectional flow—higher layers (e.g., the Platform Engine) may depend on lower layers (e.g., Domain Engine), but lower layers (e.g., Core Intelligence) remain completely detached and fundamentally unaware of the layers above. This completely obliterates circular dependencies and mitigates large-scale codebase rot.")
    
    add_h2(doc, "4.2 Technology Stack Selection")
    tstack = doc.add_table(rows=1, cols=3)
    tstack.style = 'Table Grid'
    t1_cells = tstack.rows[0].cells
    t1_cells[0].text = 'Category'
    t1_cells[1].text = 'Tool/Framework'
    t1_cells[2].text = 'Justification'
    
    tech_data = [
        ("API & Gateway", "FastAPI", "High performance async framework perfectly suited for concurrent IO-bound AI requests."),
        ("Routing & LLMs", "LiteLLM", "Unifies 100+ LLM interfaces into a single standard API, circumventing provider lock-in natively."),
        ("Vector DB", "Qdrant", "Offers advanced high-speed categorical filtering essential for executing complex RBAC queries concurrently with semantic search."),
        ("Orchestration", "LangGraph", "Permits highly structured, cyclic, and deterministic state-machine workflow definitions as opposed to open-ended conversational chains."),
        ("Persistence", "PostgreSQL / SQLModel", "Offers relational stability combined with straightforward zero-boilerplate Pydantic integrations for validation."),
    ]
    for cat, tool, just in tech_data:
        r = tstack.add_row().cells
        r[0].text = cat
        r[1].text = tool
        r[2].text = just
        
    add_p(doc, "")
    add_h2(doc, "4.3 Testing and Quality Assurance Methodology")
    add_p(doc, "Given that generative models inherently drift in their lexical behavior, our quality assurance paradigms fundamentally diverge from traditional strict equality testing. We employ an LLM-as-a-judge system alongside vast deterministic regression test suites.")
    add_p(doc, "Evaluative metrics such as RAG 'Faithfulness' (determining if the model hallucinated facts outside the explicitly provided text chunk) and 'Context Precision' (evaluating whether the vector DB search genuinely retrieved relevant semantic fragments over arbitrary noise) are strictly calculated and graphed nightly to alert maintaining engineers of systemic degradation.")
    doc.add_page_break()

    # 6. Datasets Used / Curation
    add_h1(doc, "5. Datasets Used and Curation Process")
    add_h2(doc, "5.1 Source Material Identification")
    add_p(doc, "The underlying AI mechanisms heavily rely on Retrieval-Augmented Generation (RAG). Therefore, the data pipeline quality directly dictates system authority. The platform curates unstructured data across multiple modalities: intricate markdown texts, semi-structured Excel spreadsheets, and strictly structured tabular JSON payloads extracted from historical API logs.")
    
    add_h2(doc, "5.2 The Curation and Ingestion Pipeline")
    add_p(doc, "Our automated ETL (Extract, Transform, Load) pipelines ingest enterprise domains through a highly stringent and linear processing protocol:")
    add_h3(doc, "Phase 1: Normalization & Cleansing")
    add_p(doc, "Source documents are strictly stripped of structural HTML/formatting inconsistencies. Highly tuned PII (Personally Identifiable Information) masking models aggressively scan and redact sensitive entity data (social security numbers, bank routing keys) using robust Regex hierarchies layered underneath lightweight NER (Named Entity Recognition) models.")
    
    add_h3(doc, "Phase 2: Semantic Chunking")
    add_p(doc, "Rather than indiscriminately splitting documents by arbitrary character counts (which inherently cleaves sentences in half and destroys lexical integrity), the ingestion system leverages structural chunking. It dynamically respects paragraph line breaks, markdown headers configurations, and semantic token boundaries, guaranteeing no loss of contextual fidelity within the vector representations.")
    
    add_h3(doc, "Phase 3: Vectorization and Metadata Extrapolation")
    add_p(doc, "The cleaned logical chunks are immediately mathematically embedded using high-dimensional dense embedding models (e.g., text-embedding-ada-002 paradigm). Vitally, they are appended with a granular metadata payload explicitly including tenant_id, security access_tier, the precise temporal_timestamp, and categorical grouping tags. This enables hybrid dense-sparse analytical queries during runtime.")

    add_h2(doc, "5.3 Synthetic Data Generation for Evaluation")
    add_p(doc, "To rigorously bullet-proof the platform against chaotic and undefined user inputs, Layer 3 heavily deploys 'Synthetic Data Extraction'. By systematically prompting mathematically superior models (such as flagship tier routers) to impersonate adversarial threat actors, the platform autonomously populates its Golden Datasets with thousands of highly stressful, context-free, and misdirected interactions to comprehensively train triage mechanisms.")
    doc.add_page_break()

    # 7. Multi-Tenancy Deep Dive
    add_h1(doc, "6. Supplemental Analysis: Multi-Tenancy Architecture Constraints")
    add_h2(doc, "6.1 The necessity of Strict Logical Isolation")
    add_p(doc, "A critical pillar missing in modern wrapper applications is absolute horizontal separation of customer assets. In a B2B SaaS environment running a single unified generative engine, the possibility of data pollution across clients defines maximum regulatory risk.")
    add_p(doc, "Our architecture achieves logical separation without the exorbitant cost of deploying full independent database clusters for each tenant. Utilizing Postgres Row-Level Security (RLS), every table enforcing business logic mandates a `tenant_id` foreign key. Query interfaces inherently inject the JWT-verified tenant context into the overarching SQL transaction state, ensuring cross-tenant queries mathematically evaluate to zero.")
    
    add_h2(doc, "6.2 Vector Isolation and Hybrid Sub-Filtering")
    add_p(doc, "Standard memory abstractions dump all learned knowledge into an overarching Pinecone or Weaviate index. The architecture mandates Qdrant collections separated either logically via mandatory Payload Filtering masks or, for strict-compliance enterprise tiers, physically isolated Qdrant clusters per organizational unit.")
    add_p(doc, "This guarantees that an AI Agent performing a semantic search across “historical employee incident reports” restricted to Tenant A structurally cannot retrieve vectors computed from Tenant B’s identical workflow schema. The routing layer fundamentally rejects any prompt lacking a signed contextual tenant token before embedding construction even begins.")
    doc.add_page_break()

    # 8. Load Balancing and Cloud Deployments
    add_h1(doc, "7. Supplemental Analysis: Enterprise Rollout Methodologies")
    add_h2(doc, "7.1 Infrastructure as Code (IaC) Provisioning")
    add_p(doc, "Scaling from proof-of-concept to thousands of concurrent users necessitates extreme infrastructure replicability. The Enterprise API interfaces—predominantly containerized FastAPI workers—are provisioned entirely via Terraform modules defining Virtual Private Clouds (VPCs) across AWS or Google Cloud.")
    add_p(doc, "Auto-scaling groups observe primary trigger metrics. Given that inference queues are heavily I/O bound rather than CPU-blocked (due to awaiting external HTTP responses from LLM endpoints), the scaling rules evaluate concurrency latency parameters. When connection limits saturate, elastic schedulers spawn parallel container replicas securely behind an Application Load Balancer to dilute traffic organically.")

    add_h2(doc, "7.2 State Management under Ephemeral Workers")
    add_p(doc, "In order to construct complex cyclic Agentic Workflows (Layer 2), the system implements LangGraph states. Crucially, as web pods scale dynamically up or down and connections get interrupted, the memory of an ongoing task cannot reside inside local container RAM.")
    add_p(doc, "The statefulness of these cyclical agent transactions is aggressively committed to an asynchronous Redis cluster structure using sophisticated checkpointing operations. Thus, if a worker node crashes mid-reasoning cycle, another pod instantly assumes the exact conversational and operational state checkpoint from Redis, resuming execution deterministically without impacting end-user experience.")
    doc.add_page_break()
    
    # 9. Safety and Observability
    add_h1(doc, "8. Supplemental Analysis: Total Observability and Safety Firewalls")
    add_h2(doc, "8.1 Guardrails Implementation")
    add_p(doc, "The 'Safety & Alignment' parameter within Layer 1 isn't theoretical; it executes functional Regex boundary enforcement and sentiment filtration immediately adjacent to the core intent router. Prompts containing explicit directives intended to subvert the system's operational guidelines (Prompt Injection and Jailbreak attempts) invoke immediate refusal logic.")
    add_p(doc, "This is fundamentally critical for Layer 2 security. Before an orchestration workflow signals a system to process a refund or write data to an enterprise SQL index, the transactional logic confirms that the preceding generative state underwent 'Compliance Verification'—an assertion appended to the internal object graph confirming the output violates zero domain restrictions.")

    add_h2(doc, "8.2 Unified Telemetry Parsing")
    add_p(doc, "All applications emit exhaustive, structured JSON logs via Structlog. Rather than parsing raw string aggregates, these structured telemetry payloads (comprising `user_id`, `layer_invoked`, `compute_ms`, `tokens_consumed`, `model_used`) are continuously streamed into Datadog or Arize Phoenix.")
    add_p(doc, "This allows operations teams to trace identical workflow signatures across temporal scales. If the `gpt-4-turbo` latency average spikes globally >2000ms, monitoring scripts can instantly actuate the dynamic router to default globally to local-LLM architectures while broadcasting incident degradation flags seamlessly.")
    doc.add_page_break()

    # 10. Conclusion Rewrite
    add_h1(doc, "9. Conclusion")
    add_p(doc, "The proposed Enterprise AI Platform directly confronts and neutralizes the fragile, hard-coded implementations overwhelmingly prevalent in modern superficial AI adoption. Transitioning from unstructured prompt-chains to a rigid deterministic orchestration lattice insulates organizations against the chaos intrinsic to stochastic token generation.")
    add_p(doc, "By enforcing an absolute architectural division between cognitive comprehension spaces (Layer 1) and destructive, stateful execution networks (Layer 2), enterprises guarantee data sovereignty. They eliminate provider monopoly lock-in via standardized Layer 0 proxies, ensure robust client partitioning, and establish perpetual agility that dynamically accommodates the accelerating evolution of base-model intelligence.")
    doc.add_page_break()

    file_name = "Synopsis_Enterprise_AI_Platform_V2.docx"
    doc.save(file_name)
    print(f"Document successfully created at {os.path.abspath(file_name)}")
    print(f"File size approximations: Generated unique architectural appendices for full requirement compliance.")

if __name__ == "__main__":
    main()
